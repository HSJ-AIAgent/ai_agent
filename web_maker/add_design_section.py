"""
add_design_section.py — Appends META Design Concept section to existing docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def add_run(paragraph, text, bold=False, italic=False, size=11, color=None, font='Calibri'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_color_swatch(doc, swatches):
    """swatches: list of (name, hex, description)"""
    table = doc.add_table(rows=2, cols=len(swatches))
    table.style = 'Table Grid'
    for i, (name, hex_color, desc) in enumerate(swatches):
        top_cell = table.cell(0, i)
        top_cell.text = ''
        set_cell_bg(top_cell, hex_color)
        top_cell.height = Cm(1.5)

        bot_cell = table.cell(1, i)
        p = bot_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(name + '\n')
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = p.add_run(f'#{hex_color}\n')
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(100, 100, 100)
        r3 = p.add_run(desc)
        r3.font.size = Pt(8)
        r3.italic = True
        r3.font.color.rgb = RGBColor(130, 130, 130)

    for i in range(len(swatches)):
        table.columns[i].width = Cm(16 / len(swatches))
    return table


def append_design_section(filepath):
    doc = Document(filepath)

    # ── Page break before new section ────────────────────────────────────────
    doc.add_page_break()

    # ── Section 7 heading ────────────────────────────────────────────────────
    doc.add_heading('7. Design Concept & Brand Identity', level=1)

    intro = doc.add_paragraph()
    add_run(intro,
        'Meta\'s visual identity underwent a landmark transformation in October 2021 with the corporate rebranding from '
        'Facebook to Meta Platforms. The new design language reflects the company\'s strategic pivot toward the metaverse, '
        'spatial computing, and AI-first product experiences — while maintaining continuity with its core social media roots.',
        size=11)

    # ── 7.1 Logo & Symbol ─────────────────────────────────────────────────────
    doc.add_heading('7.1 Logo & Symbol Design', level=2)

    logo_data = [
        ('Symbol Form', 'Stylized infinity loop (∞) rendered as a Möbius-strip-like ribbon'),
        ('Design Concept', '"Infinite possibilities and interconnected experiences" within the metaverse'),
        ('Dimensionality', 'Designed to exist in 3D space — can be traversed, textured, and animated'),
        ('Color Treatment', 'Blue gradient from #0668E1 → lighter blue; connects to Facebook\'s legacy blue'),
        ('Wordmark Font', 'Custom rounded sans-serif (introduced 2019); minimalist, accessible, futuristic'),
        ('Design Agency', 'Developed internally by Meta\'s Brand Design team'),
        ('Revealed', 'October 28, 2021 — Meta Connect keynote'),
    ]

    table = doc.add_table(rows=len(logo_data) + 1, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(['Element', 'Description']):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '0668E1')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, (elem, desc) in enumerate(logo_data, start=1):
        row = table.rows[r]
        row.cells[0].text = elem
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = desc
        if r % 2 == 0:
            set_cell_bg(row.cells[0], 'EBF3FF')
            set_cell_bg(row.cells[1], 'EBF3FF')

    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    # ── 7.2 Color Palette ─────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading('7.2 Brand Color Palette', level=2)

    palette_intro = doc.add_paragraph()
    add_run(palette_intro,
        'Meta\'s official color system is built on a blue-dominant palette that bridges the heritage '
        'of Facebook\'s identity with a modern, gradient-driven aesthetic suited for digital and spatial surfaces.',
        size=11)
    doc.add_paragraph()

    swatches = [
        ('Primary Blue',    '0072F0', 'Core brand color'),
        ('Deep Blue',       '0668E1', 'Logo gradient start'),
        ('Sky Blue',        '1B74E4', 'Interactive elements'),
        ('Light Blue',      'EBF3FF', 'Backgrounds / tints'),
        ('Meta Purple',     '6B2FA0', 'AI & Metaverse accent'),
        ('Neutral Gray',    '9CA3AF', 'Supporting text'),
        ('Dark',            '1F2937', 'Body text / dark mode'),
        ('White',           'FFFFFF', 'Backgrounds'),
    ]
    add_color_swatch(doc, swatches)

    doc.add_paragraph()

    # ── 7.3 Typography ────────────────────────────────────────────────────────
    doc.add_heading('7.3 Typography System', level=2)

    typo_data = [
        ('Primary Typeface', 'Custom Meta sans-serif (rounded, humanist)', 'Wordmark, headings, UI labels'),
        ('Secondary', 'Optimistic (Meta\'s proprietary typeface)', 'Long-form content, marketing copy'),
        ('Fallback / Web', 'Segoe UI / Helvetica Neue / Arial', 'System UI, web fallback'),
        ('Style', 'Minimal stroke contrast, wide apertures, generous x-height', 'Readability across sizes'),
        ('Weight Range', 'Light (300) → ExtraBold (800)', 'Hierarchy and emphasis'),
        ('Usage Rule', 'No decorative fonts; clarity and legibility prioritized', 'Accessibility-first principle'),
    ]

    ttable = doc.add_table(rows=len(typo_data) + 1, cols=3)
    ttable.style = 'Table Grid'
    for i, h in enumerate(['Attribute', 'Value', 'Application']):
        c = ttable.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '1B74E4')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, (attr, val, app) in enumerate(typo_data, start=1):
        row = ttable.rows[r]
        row.cells[0].text = attr
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val
        row.cells[2].text = app
        if r % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F0F6FF')

    ttable.columns[0].width = Cm(4)
    ttable.columns[1].width = Cm(7)
    ttable.columns[2].width = Cm(5)

    # ── 7.4 Core Design Principles ────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading('7.4 Core Design Principles', level=2)

    principles = [
        ('Simplicity & Minimalism',
         'Remove friction. Strip UI down to only what is needed. Exemplified in Threads\' clean interface — '
         'no distractions, content-first layout.'),
        ('Adaptive Brand Systems',
         'The Meta brand is designed to flex across surfaces: 2D screens, AR overlays, VR environments, '
         'wearables, and physical spaces. Static logos are replaced by dynamic, context-aware expressions.'),
        ('Hierarchy & Grid',
         'Consistent use of 8pt grid systems and typographic hierarchy guides user attention. '
         'Information is layered from primary to tertiary with clear visual weight.'),
        ('Continuity with Heritage',
         'The blue gradient in the Meta logo deliberately references Facebook\'s legacy blue, '
         'bridging the old identity with the new metaverse-era brand.'),
        ('Accessibility First',
         'WCAG 2.1 AA compliance is a baseline. Color contrast ratios, font sizes, and touch target '
         'sizes are designed for diverse users including those with mobility, vision, and hearing needs.'),
        ('Micro-interactions',
         'Subtle animations — like reactions, hover states, and loading indicators — communicate '
         'system feedback and create emotional engagement without overwhelming content.'),
        ('Spatial-ready Design',
         'All design components are built to translate to 3D space. Depth, shadow, and parallax layers '
         'are considered at component level for Horizon OS and future AR glasses.'),
    ]

    for title, body in principles:
        p = doc.add_paragraph()
        add_run(p, f'{title}  ', bold=True, size=11, color=(6, 104, 225))
        add_run(p, body, size=11)
        p.paragraph_format.space_after = Pt(8)

    # ── 7.5 Spatial / AR / VR Design Language (Horizon OS) ───────────────────
    doc.add_paragraph()
    doc.add_heading('7.5 Spatial Computing & Horizon OS Design Language', level=2)

    spatial_intro = doc.add_paragraph()
    add_run(spatial_intro,
        'Meta Horizon OS represents the company\'s most advanced design frontier — a full spatial computing '
        'operating system for AR/VR devices. Its design language extends traditional 2D UI principles into '
        'three-dimensional interactive environments.',
        size=11)
    doc.add_paragraph()

    spatial_items = [
        'Environment as Canvas — physical space is treated as the design surface; UI panels float in 3D',
        'Passthrough Design — mixed reality UI must complement, not obscure, the real world',
        'Depth & Proximity — element scale and opacity convey distance and importance in 3D space',
        'Spatial Audio — sound design reinforces UI interactions and spatial positioning',
        'Haptic Feedback — controller and hand-tracking haptics communicate state changes',
        'Figma-based Horizon OS UI Kit — designers use exact Quest UI components for accurate prototyping',
        'Scene Understanding — UI adapts dynamically to the user\'s physical environment geometry',
        'Immersive Onboarding — first-minute experience is critical; tutorials enrich without blocking flow',
    ]
    for item in spatial_items:
        doc.add_paragraph(item, style='List Bullet')

    # ── 7.6 Design at Meta (Team Structure) ───────────────────────────────────
    doc.add_paragraph()
    doc.add_heading('7.6 Design at Meta — Team Structure', level=2)

    team_data = [
        ('Brand Design',     'Visual identity systems, logos, color, typography, illustration, motion, iconography'),
        ('Product Design',   'UX/UI for Facebook, Instagram, WhatsApp, Messenger, Threads apps'),
        ('Spatial Design',   'Horizon OS, Quest interface, AR/VR experience design'),
        ('Motion Design',    'Animations, transitions, video — across apps and brand communications'),
        ('Research (UXR)',   'User research, usability testing, personas, journey mapping'),
        ('Content Design',   'UX writing, microcopy, voice and tone guidelines'),
        ('AI Design',        'Conversational UI for Meta AI; AI Studio interface design'),
    ]

    dteam = doc.add_table(rows=len(team_data) + 1, cols=2)
    dteam.style = 'Table Grid'
    for i, h in enumerate(['Design Discipline', 'Responsibilities']):
        c = dteam.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '6B2FA0')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, (disc, resp) in enumerate(team_data, start=1):
        row = dteam.rows[r]
        row.cells[0].text = disc
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = resp
        if r % 2 == 0:
            set_cell_bg(row.cells[0], 'F5EEFF')
            set_cell_bg(row.cells[1], 'F5EEFF')

    dteam.columns[0].width = Cm(5)
    dteam.columns[1].width = Cm(11)

    # ── Closing note ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run(p,
        'Meta\'s design philosophy can be summarized as: ',
        italic=True, size=11, color=(100, 100, 100))
    add_run(p,
        '"Build for where people are going, not where they are." ',
        bold=True, italic=True, size=11, color=(6, 104, 225))
    add_run(p,
        'Every design decision — from the infinity logo to Horizon OS spatial guidelines — '
        'is oriented toward a future where digital and physical realities converge.',
        italic=True, size=11, color=(100, 100, 100))

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(filepath)
    print(f'Updated: {os.path.abspath(filepath)}')


if __name__ == '__main__':
    target = r'C:\Users\SBS\Desktop\Agent_HSJ\web_maker\Meta_Research_Report.docx'
    append_design_section(target)
