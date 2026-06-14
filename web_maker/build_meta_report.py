"""
build_meta_report.py — META Research Report generator
Saves to: web_maker/Meta_Research_Report.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    for tag in ['begin', None, 'end']:
        if tag:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        else:
            el = OxmlElement('w:instrText')
            el.text = 'PAGE'
        run._r.append(el)


def add_run(paragraph, text, bold=False, italic=False, size=11, color=None, font='Calibri'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def build():
    doc = Document()

    # ── Page layout (A4) ─────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Header ───────────────────────────────────────────────────────────────
    hdr = section.header.paragraphs[0]
    hdr.clear()
    add_run(hdr, 'META Platforms, Inc. — Research Report', bold=True, size=9, color=(0, 114, 240))
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ── Footer ───────────────────────────────────────────────────────────────
    ftr = section.footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(ftr, 'Page ', size=9)
    add_page_number(ftr)
    add_run(ftr, '  |  Confidential — For Internal Use Only', size=9, italic=True, color=(120, 120, 120))

    # ── Cover ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title_p, 'META Platforms, Inc.', bold=True, size=28, color=(0, 114, 240))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub_p, 'Comprehensive Research Report', bold=True, size=16, color=(60, 60, 60))

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(date_p, 'June 2026  |  Prepared by AI Research Agent', italic=True, size=11, color=(120, 120, 120))

    doc.add_paragraph()

    # Divider line
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0072F0')
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_page_break()

    # ── 1. Company Overview ──────────────────────────────────────────────────
    doc.add_heading('1. Company Overview', level=1)

    overview_data = [
        ('Full Name', 'Meta Platforms, Inc.'),
        ('Founded', '2004 (as TheFacebook, Harvard University)'),
        ('Headquarters', 'Menlo Park, California, USA'),
        ('CEO', 'Mark Zuckerberg (Co-founder)'),
        ('Stock Ticker', 'META (NASDAQ)'),
        ('Employees', '77,986 (as of March 2026)'),
        ('Rebranding', 'Facebook → Meta Platforms (October 2021)'),
    ]

    table = doc.add_table(rows=len(overview_data) + 1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(['Attribute', 'Details']):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], '0072F0')
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, (attr, val) in enumerate(overview_data, start=1):
        row = table.rows[r]
        row.cells[0].text = attr
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val
        if r % 2 == 0:
            set_cell_bg(row.cells[0], 'EBF3FF')
            set_cell_bg(row.cells[1], 'EBF3FF')

    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run(p, 'History & Milestones', bold=True, size=12)

    milestones = [
        '2004 — TheFacebook launched at Harvard by Mark Zuckerberg',
        '2006 — Opened to the public (anyone 13+)',
        '2012 — IPO on NASDAQ; acquired Instagram for ~$1B',
        '2014 — Acquired WhatsApp for $19B; acquired Oculus VR for $2B',
        '2016 — Launched Marketplace and Live video features',
        '2021 — Rebranded to Meta Platforms; launched metaverse initiative',
        '2023 — Launched Threads (Twitter/X competitor)',
        '2025 — LLaMA 4 released; Instagram reached 3B MAU; WhatsApp 3B+ users',
        '2026 — Meta Superintelligence Labs established; Q1 Family DAP: 3.56B',
    ]
    for m in milestones:
        doc.add_paragraph(m, style='List Bullet')

    # ── 2. Core Products ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('2. Core Products & Services', level=1)

    products = [
        ('Facebook', '3.07B', 'Social networking, News Feed, Groups, Marketplace, Live'),
        ('Instagram', '3.0B', 'Photo/video sharing, Reels, Stories, Shopping, DMs'),
        ('WhatsApp', '3.0B+', 'Encrypted messaging, Voice/Video calls, Business API'),
        ('Messenger', '~1B+', 'Chat, video calls, integrated with Facebook'),
        ('Threads', 'Growing', 'Text-based social network, Twitter/X alternative'),
        ('Meta Quest', 'VR/AR HW', 'Virtual Reality headsets (Quest 3), mixed reality'),
        ('Meta AI', 'Integrated', 'AI assistant across all Meta apps and Ray-Ban glasses'),
        ('Ray-Ban Smart Glasses', 'Wearable', 'AI-powered glasses with Meta AI built-in'),
    ]

    ptable = doc.add_table(rows=len(products) + 1, cols=3)
    ptable.style = 'Table Grid'
    for i, h in enumerate(['Product', 'MAU / Scale', 'Key Features']):
        c = ptable.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '1B74E4')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, (prod, mau, feat) in enumerate(products, start=1):
        row = ptable.rows[r]
        row.cells[0].text = prod
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = mau
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = feat
        if r % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F0F6FF')

    ptable.columns[0].width = Cm(4)
    ptable.columns[1].width = Cm(3)
    ptable.columns[2].width = Cm(9)

    # ── 3. Financial Performance ─────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('3. Financial Performance', level=1)

    p = doc.add_paragraph()
    add_run(p, 'Meta achieved record financial results in 2025, with revenue surpassing $200 billion for the first time. '
               'The company continues to demonstrate strong profitability while making massive investments in AI infrastructure for 2026.', size=11)

    doc.add_paragraph()
    doc.add_heading('Annual Revenue & Profit (2023–2026E)', level=2)

    fin_data = [
        ('FY 2023', '$134.9B', '$39.1B', '$23.2B', '29%'),
        ('FY 2024', '$164.5B', '$58.1B', '$62.4B', '35%'),
        ('FY 2025', '$200.97B', 'N/A', '$72.4B est.', '36%'),
        ('Q1 2026', '$42.3B est.', 'N/A', 'N/A', '~37%'),
    ]

    ftable = doc.add_table(rows=len(fin_data) + 1, cols=5)
    ftable.style = 'Table Grid'
    fin_headers = ['Period', 'Revenue', 'Operating Income', 'Net Income', 'Op. Margin']
    for i, h in enumerate(fin_headers):
        c = ftable.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '0A4A8A')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, row_data in enumerate(fin_data, start=1):
        row = ftable.rows[r]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'E8F0FA')

    doc.add_paragraph()
    doc.add_heading('2026 Capital Expenditure Guidance', level=2)
    capex_items = [
        'Total 2026 CapEx guidance: $125B – $145B (AI infrastructure focus)',
        '2026 total expenses expected: $162B – $169B',
        'Q2 2026 revenue guidance: $58B – $61B',
        'Operating income expected to exceed 2025 levels',
        'Primary CapEx driver: data centers and AI compute hardware',
    ]
    for item in capex_items:
        doc.add_paragraph(item, style='List Bullet')

    # ── 4. AI Initiatives ────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('4. AI Initiatives & Strategy', level=1)

    p = doc.add_paragraph()
    add_run(p, 'Meta is aggressively positioning itself as an AI-first company. '
               'Under the "Superintelligence Era" strategy, AI is being embedded across all products '
               'and the company has established Meta Superintelligence Labs as a dedicated research division.', size=11)

    doc.add_paragraph()
    doc.add_heading('LLaMA Model Family', level=2)

    llama_data = [
        ('LLaMA 1', 'Feb 2023', 'First open-source LLM release; research-focused'),
        ('LLaMA 2', 'Jul 2023', 'Open commercial use; 7B–70B parameters'),
        ('LLaMA 3', 'Apr 2024', 'Multilingual; 8B–405B params; major capability leap'),
        ('LLaMA 4', 'Apr 2025', 'Mixture of Experts (MoE) architecture; Scout, Maverick, Behemoth variants'),
        ('LLaMA 4.5 / 4.X', 'Late 2026 (planned)', 'Next-gen model; advanced reasoning and multimodal'),
    ]

    ltable = doc.add_table(rows=len(llama_data) + 1, cols=3)
    ltable.style = 'Table Grid'
    for i, h in enumerate(['Model', 'Release', 'Key Features']):
        c = ltable.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '6B2FA0')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, (model, release, feat) in enumerate(llama_data, start=1):
        row = ltable.rows[r]
        row.cells[0].text = model
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = release
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = feat
        if r % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F5EEFF')

    ltable.columns[0].width = Cm(4)
    ltable.columns[1].width = Cm(4)
    ltable.columns[2].width = Cm(8)

    doc.add_paragraph()
    doc.add_heading('Key AI Products & Initiatives', level=2)

    ai_items = [
        'Meta AI Assistant — integrated across Facebook, Instagram, WhatsApp, Messenger, and Ray-Ban glasses',
        'Meta AI Studio — platform for creating AI characters and chatbots (targeting #1 AI character destination)',
        'Meta Movie Gen — AI video generation and editing research models',
        'Meta Superintelligence Labs — dedicated division for frontier AI research',
        'Llama for Startups (May 2025) — ecosystem program creating Llama-native companies',
        'Sub-Saharan Africa AI Initiative — open-source AI development in Nigeria, Kenya, Senegal, South Africa',
        'AI-powered advertising tools — automated ad creation, targeting, and optimization',
        'Reality Labs — AR/VR research combining AI with spatial computing',
    ]
    for item in ai_items:
        doc.add_paragraph(item, style='List Bullet')

    # ── 5. Competitive Position ──────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('5. Market Position & Competitive Landscape', level=1)

    comp_data = [
        ('Meta Platforms', 'Social Media / AI', '3.56B DAP', '$200.97B (2025)', 'Dominant social + AI pivot'),
        ('Alphabet (Google)', 'Search / AI / Cloud', '~2B+ MAU (YouTube)', '$350B+', 'Gemini AI, search monopoly'),
        ('Microsoft', 'Cloud / AI / Enterprise', 'Varied', '$245B+', 'OpenAI partnership, Azure AI'),
        ('Apple', 'Hardware / Ecosystem', '~2B devices', '$391B+', 'Apple Intelligence, hardware moat'),
        ('ByteDance (TikTok)', 'Short Video', '1.7B+ MAU', '$120B est.', 'Competing with Reels'),
        ('X (Twitter)', 'Social / AI (Grok)', '~250M MAU', 'Private', 'Declining but niche influence'),
    ]

    ctable = doc.add_table(rows=len(comp_data) + 1, cols=5)
    ctable.style = 'Table Grid'
    for i, h in enumerate(['Company', 'Sector', 'Scale', 'Revenue', 'Key Differentiator']):
        c = ctable.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '2D4A7A')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r, row_data in enumerate(comp_data, start=1):
        row = ctable.rows[r]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
        if r == 1:
            for c in row.cells:
                set_cell_bg(c, 'D6E4FF')
        elif r % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, 'F5F5F5')

    # ── 6. Summary ───────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('6. Executive Summary', level=1)

    summary_points = [
        ('User Scale', 'Meta\'s family of apps serves 3.56 billion daily active people — unmatched in the industry.'),
        ('Financial Strength', '$200.97B revenue in 2025 (22% YoY growth); massive $125–145B CapEx planned for 2026 AI infrastructure.'),
        ('AI Leadership', 'Open-source LLaMA strategy creates a broad AI ecosystem; Meta AI integrated across all touchpoints.'),
        ('Product Diversification', 'Multiple 3B+ MAU platforms (Facebook, Instagram, WhatsApp) reduce dependency on any single app.'),
        ('Hardware Ambition', 'Reality Labs + Ray-Ban smart glasses signal long-term bet on spatial computing and AI wearables.'),
        ('Risk Factors', 'Heavy CapEx commitments, regulatory pressure on data privacy, and intense competition from ByteDance/TikTok.'),
    ]

    for title, body in summary_points:
        p = doc.add_paragraph()
        add_run(p, f'{title}: ', bold=True, size=11, color=(0, 72, 160))
        add_run(p, body, size=11)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '— End of Report —', italic=True, size=10, color=(150, 150, 150))

    # ── Save ─────────────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Meta_Research_Report.docx')
    doc.save(output_path)
    print(f'Saved: {output_path}')
    return output_path


if __name__ == '__main__':
    build()
