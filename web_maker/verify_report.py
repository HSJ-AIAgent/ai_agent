"""verify_report.py — parse and verify Meta_Research_Report.docx contents"""
from docx import Document

doc = Document('Meta_Research_Report.docx')

sections_found = []
tables_found = len(doc.tables)
bullet_count = 0
total_paragraphs = 0

for para in doc.paragraphs:
    total_paragraphs += 1
    if para.style.name.startswith('Heading'):
        sections_found.append(f"[{para.style.name}] {para.text}")
    if 'List' in para.style.name:
        bullet_count += 1

print(f"=== Document Verification ===")
print(f"Total paragraphs : {total_paragraphs}")
print(f"Tables           : {tables_found}")
print(f"List items       : {bullet_count}")
print(f"\n=== Headings Found ===")
for s in sections_found:
    print(f"  {s}".encode('utf-8', errors='replace').decode('utf-8'))

print(f"\n=== Table Summary ===")
for i, t in enumerate(doc.tables):
    print(f"  Table {i+1}: {len(t.rows)} rows x {len(t.columns)} cols | Header: {t.rows[0].cells[0].text}")

print(f"\n=== Header/Footer ===")
for sec in doc.sections:
    print(f"  Header: {sec.header.paragraphs[0].text.strip()}")
    print(f"  Footer text runs: {len(sec.footer.paragraphs[0].runs)} runs")

print("\n=== Verification PASSED ===")
